from langchain_openai import ChatOpenAI
from dotenv import load_dotenv
import os
from helper_functions import parse_response, connect_db
from langchain_core.runnables import RunnableLambda
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
import pickle
from helper_functions import build_prompt, parse_response
from langchain_core.runnables import RunnablePassthrough, RunnableLambda
from langchain.retrievers import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import LLMChainExtractor
import base64
from langchain_core.messages import SystemMessage, HumanMessage
from langchain_openai import ChatOpenAI
from base64 import b64decode
import tiktoken
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
from fpdf import FPDF
import asyncio
from docx import Document
from models import Reports
from io import BytesIO
from docx.shared import Inches
import requests
import json
from pathlib import Path
import re

# Load API key
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
openai_api_key = os.getenv("OPENROUTER_API_KEY")

genai.configure(api_key=api_key)


doc = Document()


#This is not as good as 2.5 pro, but use this in case the gemini tokens that karthik gave run out
'''
model = ChatOpenAI(
    temperature=0.5,
    model_name="google/gemini-2.5-flash-preview-05-20",  
    base_url="https://openrouter.ai/api/v1",
    api_key=openai_api_key
)
'''

#better chat model for report generation
model = ChatGoogleGenerativeAI(
    model="gemini-2.5-pro", 
    google_api_key=api_key,
    temperature=0.3
)


#refresh token to access image url for microsoft graph api
def refresh_access_token(token, token_path): 
    print("Refreshing access token...")
    scopes = ['User.Read']

    tenant_id = "4532a721-a36d-41e1-8367-80fd926092a2"
    client_id = "4d7dd08e-dc2d-4334-961f-93506b22e435"
    client_secret = "eRJ8Q~DzN7Rxkzyihq0mQedZuXFYzqFXETz4CaEB"
    refresh_token = token['refresh_token']

    token_url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"

    # Data sent to Microsoft to get a new access token
    data = {
        "grant_type": "refresh_token",      # Tells MS you want to refresh a token
        "client_id": client_id,              # Your app ID
        "refresh_token": refresh_token,      # The long-lived token you got earlier
        "scope": "https://graph.microsoft.com/.default"  # Permissions you want
    }

    response = requests.post(token_url, data)

    if response.ok: 
        new_token = response.json()
        print("Token successfully refreshed!")

        with open(token_path, "w") as f: 
            json.dump(new_token, f, indent=2)

        return new_token
    
    else: 
        raise Exception("Failed to refresh token", response.text)



# Access tokens that give authorization to use OneNote API
token_path = Path.home() / ".credentials" / "onenote_graph_token.json"
with open(token_path, "r") as f:
    token = json.load(f)

headers = {
    "Authorization": f"Bearer {token['access_token']}"
}





def get_all_summaries(retriever):

    # Directly fetch all vectors from the Chroma store
    return [doc for doc in retriever.vectorstore.get()['documents']]

async def generate_introduction_agent(retriever):
    summaries = get_all_summaries(retriever)
    context = "\n".join(summaries)

    prompt_template = """
        You are a scientific writer. Write the **Introduction** section of a technical report.

        Describe the motivation for the project, the problem context, and what was attempted.

        Use **only** the following context. Do **not** fabricate or infer external details.
        Write clearly and professionally for a technical audience.

        ---
        Context:
        {context}
        ---

        **Example Output:**

        Introduction:  
        Allopregnanolone is a neuroactive steroid under investigation for its potential in treating neurological and neurodegenerative conditions. Its low aqueous solubility presents a major challenge in formulating it for parenteral or transdermal delivery. Several patents currently exist covering SEDDS and micellar formulations, necessitating novel approaches that avoid infringement. This study was initiated to systematically evaluate a broad range of excipients and solvent systems, using both empirical and computational tools to guide formulation development.
    """

    prompt = ChatPromptTemplate.from_template(prompt_template)
    chain = prompt | model | StrOutputParser()
    return await chain.ainvoke({"context": context})

async def generate_summary_agent(retriever):
    summaries = get_all_summaries(retriever)
    context = "\n".join(summaries)

    prompt_template = """
        You are a scientific writer. Write the **Summary** of a technical report.

        Provide a concise, high-level overview of the key findings, results, and outcomes.

        Use **only** the provided context. Do **not** add external information.
        Write in a professional and concise tone suitable for a technical audience.

        ---
        Context:
        {context}
        ---

        **Example Output:**

        Summary:  
        The goal of this project was to develop a formulation of allopregnanolone (Allo) with a target solubility of 6 mg/mL in saline, suitable for intravenous (IV) or transdermal administration. Given the hydrophobic nature of Allo and limited water solubility, solvent screening, excipient selection, and emulsion development studies were conducted. Initial efforts focused on evaluating various solvents and surfactants. The excipient testing was guided by both empirical screening and further computational predictions. Emulsions using Span 20 and Tween 20 showed potential for use, albeit with limited formulation stability.
    """

    prompt = ChatPromptTemplate.from_template(prompt_template)
    chain = prompt | model | StrOutputParser()
    return await chain.ainvoke({"context": context})


async def generate_objectives_agent(retriever):
    summaries = get_all_summaries(retriever)
    context = "\n".join(summaries)

    prompt_template = """
        You are a scientific writer. List the **Objectives** of the project based on the context.

        Write 3–5 numbered objectives using clear, concise, and goal-oriented language.

        Do not fabricate or use external info — only what's in the context.

        ---
        Context:
        {context}
        ---

        **Example Output:**

        Objectives:  
        The project had the following objectives:
          1. To develop a 6 mg/mL formulation of Allo in saline.  
          2. To evaluate physical and chemical stability over 1–3 months.  
          3. To ensure the formulation avoids infringement of existing patents.
    """

    prompt = ChatPromptTemplate.from_template(prompt_template)
    chain = prompt | model | StrOutputParser()
    return await chain.ainvoke({"context": context})


async def generate_methodology(retriever, project_name): 

    #want to retrieve the raw documents instead of summaries and feed that as context to llm

    #retrieve the unstructured data from pickle, it is in caches/all_content[section_name]
    if os.path.exists("./caches/all_content_cache.pkl"): 
        with open("./caches/all_content_cache.pkl", "rb") as f: 
            all_content = pickle.load(f)

    project_raw_documents = all_content[project_name]

    def is_methodology_related(element): 
        text = str(element).lower()

        if any(h in text for h in ["# method", "## method", "# methodology", "## methodology"]):
            return True

        # Fallback: keyword-based match (if section headers aren't included)
        keywords = [
            "protocol", "procedure", "method", "solubility", "formulation", "excipient",
            "hplc", "mixing", "optimization", "characterization", "spectrophotometry",
            "turbidity", "microscopy", "emulsion", "computational prediction"
        ]

        return any(k in text for k in keywords)
    
    filtered_elements = [str(e) for e in project_raw_documents if is_methodology_related(e)]


    prompt_template = """
    You are an expert scientific writer.

    Your task is to write the **Methodology** section of a technical report using only the provided context.

    Organize the section into relevant **subsections** with appropriate titles. Write in clear, professional **paragraph form** — avoid bullet points or numbered lists unless describing step-by-step procedures.

    Always format any tables in markdown syntax using pipes (`|`) and dashes (`-`). Do not use HTML or other formats for tables.

    Each subsection should include:
    - Materials or compounds used (e.g., solvents, excipients)
    - Procedures performed
    - Equipment, techniques, and any measured parameters
    - Optional: brief rationale for the approach
    - References to figures and tables inline using `[TABLE_X]`, `[FIGURE_X]` format

    Immediately after each figure or table reference, include:

        - The **caption** using the format:  
        **Figure X: Description of the figure content**  

        - On the line *immediately before* the caption, include **only** the raw image URL in this format (do not wrap it in tags or brackets or parentheses):  
        `Image URL: https://graph.microsoft.com/v1.0/...`

    IMPORTANT: For each figure, include both a reference and the corresponding image using a provided link (`src`). Only use images where appropriate and useful.

    Use a formal and concise scientific tone. Do not speculate or fabricate.

    If the context includes tables or figures, mention them naturally within the paragraph text by inserting inline placeholders like `[TABLE_1]` or `[FIGURE_1]` at the exact point of reference.

    Immediately after the paragraph that references a table or figure, include its caption on the next line in this exact format:

    **Table 1: Description of the data in the table**  
    **Figure 1: Description of the figure content**

    Number tables and figures sequentially throughout the entire Methodology section (i.e., TABLE_1, TABLE_2, FIGURE_1, FIGURE_2, etc.).

    Do not separate placeholders and captions with extra blank lines; keep them directly below the paragraph referencing them, so they appear interwoven in the text.

    Ensure each placeholder and caption pair corresponds clearly to unique tables or figures in the context.

    ---

    Here is an example of a high-quality Methodology section written in paragraph form:

    ### Solvent Solubility Screening
    Allopregnanolone was evaluated for solubility in a range of pharmaceutically acceptable solvents. Drug stocks were prepared at increasing concentrations ranging from 1 to 200 mg/mL in each solvent. These mixtures were vortexed for 1–2 minutes and allowed to equilibrate for 24 hours at room temperature. Preliminary solubility was assessed visually, followed by filtration through 0.22 μm PTFE filters and analysis using UV-vis spectrophotometry at 286 nm. Standard curves were constructed to determine saturation concentrations.

    ---

    Now write a new Methodology section using the context below:

    {context}

    ---
    Respond in this format:

    Methodology:

    ### <Subsection Title 1>
    <Your paragraph-form explanation>

    **Image URL**: https://graph.microsoft.com/v1.0/siteCollections/...
    **Figure 1: Burst release trends across polymer concentrations**  

    ### <Subsection Title 2>
    <Another paragraph-form explanation>
    """



    prompt = ChatPromptTemplate.from_template(prompt_template)

    chain = prompt | model | StrOutputParser()

    return await chain.ainvoke({"context": filtered_elements})

async def generate_results(retriever, project_name): 

    #want to retrieve the raw documents instead of summaries and feed that as context to llm

    #retrieve the unstructured data from pickle, it is in caches/all_content[section_name]
    if os.path.exists("./caches/all_content_cache.pkl"): 
        with open("./caches/all_content_cache.pkl", "rb") as f: 
            all_content = pickle.load(f)

    project_raw_documents = all_content[project_name]


    def is_results_related(element): 
        text = str(element).lower()

        if any(h in text for h in ["# result", "## result", "# results", "## results", "# data", "## data"]):
            return True

        # Fallback: keyword-based match (if section headers aren't included)
        keywords = [
            "result", "results", "finding", "findings", "observation", "observations",
            "data", "measurement", "measurements", "analysis", "analyses",
            "effect", "effects", "comparison", "yield", "concentration",
            "precipitate", "solubility", "turbidity", "microscopy", "hplc",
            "image", "imaging", "spectrophotometry", "table", "figure"
        ]

        return any(k in text for k in keywords)
    
    filtered_elements = [str(e) for e in project_raw_documents if is_results_related(e)]




    prompt_template = """
    You are an expert scientific writer.

    Your task is to write the **Results** section of a technical report using only the provided context.

    Organize the section into relevant **subsections** that reflect different experiments or cohorts. Write in clear, professional **paragraph form**. Avoid interpretation or speculation — stick strictly to the observed data.

    Always format any tables in markdown syntax using pipes (`|`) and dashes (`-`). Do not use HTML or other formats for tables.

    Each subsection should include:
    - Description of what was tested or measured
    - Quantitative results (e.g., concentrations, timepoints)
    - Visual or physical observations (e.g., precipitate formation, emulsion behavior)
    - References to figures and tables inline using `[TABLE_X]`, `[FIGURE_X]` format
    Immediately after each figure or table reference, include:

        - The **caption** using the format:  
        **Figure X: Description of the figure content**  

        - On the line *immediately before* the caption, include **only** the raw image URL in this format (do not wrap it in tags or brackets or parentheses):  
        `Image URL: https://graph.microsoft.com/v1.0/...`

        Keep it in that format ONLY. No extra text, brackets, tags, parentheses.

    IMPORTANT: For each figure, include both a reference and the corresponding image a provided link (`src`). Only use images where appropriate and useful.


    Maintain proper sequential numbering throughout the section (e.g., Table 4, Figure 2, etc.).

    Use a formal and concise scientific tone. Do not speculate or interpret the meaning of the results.

    When appropriate, you may include a short **Most Impactful Features** bullet list **at the end of a subsection**. These should:
    - Highlight particularly strong or notable trends
    - Focus on how specific variables influenced outcomes
    - Be phrased as short bullets like this:
        • <Variable> — <observed effect>. <Mechanistic reasoning based on the data>.

    Only include this bullet list when the data strongly supports such insight, and limit to 1–4 bullets max. Do not add bullets if the paragraph already fully explains the trend.

    ---

    Here is an example of a high-quality Results subsection:

    ### First-Round Excipient Screening

    Early methodology development focused on the development of an optimal method to
    combine the drug and excipients and subsequently remove the solvent to check the Allo
    solubility in PBS. Differential solubility of excipients in solvents mandated the use of
    multiple solvents in the study. Using different solvents had an impact on the solubility of
    the drug in combination with a specific excipient. DCM was found to be the best solvent
    for Allo screening.

    ---

    ---

    Here is an example of a high-quality Most Impactful Features section:

    ### Most Impactful Features

    • Drug (mg/mL): –0.22 — Higher drug concentration reduces burst release. This could be due to self-association or aggregation at high loading, resulting in poorer initial release.  
    • Additive (%) (PEG-3000): +0.34 — Higher PEG-3000 content increases burst release. PEG-3000, being hydrophilic, enhances matrix permeability and enables more rapid drug diffusion during early release.  
    • PLGA (mg/mL): –0.31 — Higher PLGA content lowers burst. A denser, more concentrated polymer matrix restricts rapid initial diffusion of drug.  
    • PLGA IV (viscosity): –0.39 — Higher viscosity (IV) reduces burst release. This is consistent with denser networks (higher molecular weight) limiting the initial release phase.

    ---

    Now write a Results section using the context and images below. For each figure, include both a reference and the corresponding image using either base64 or a provided link (`src`). Only use images where appropriate and useful.

    {context}

    ---

    Respond in this format:

    Results:

    ### <Subsection Title 1>
    <Paragraph-form results with inline table/figure references and captions>
    
    **Image URL**: https://graph.microsoft.com/v1.0/siteCollections/...
    **Figure 1: Burst release trends across polymer concentrations**  

    (Optional)  
    **Most Impactful Features:**  
    • <Your bullet>  
    • <Your bullet>  

    ### <Subsection Title 2>
    <More results...>
    """


    prompt = ChatPromptTemplate.from_template(prompt_template)


    chain = prompt | model | StrOutputParser()

    return await chain.ainvoke({"context": filtered_elements})


async def generate_conclusion(retriever, project_name): 

    #want to retrieve the raw documents instead of summaries and feed that as context to llm

    #retrieve the unstructured data from pickle, it is in caches/all_content[section_name]
    if os.path.exists("./caches/all_content_cache.pkl"): 
        with open("./caches/all_content_cache.pkl", "rb") as f: 
            all_content = pickle.load(f)

    project_raw_documents = all_content[project_name]

    def is_results_related(element): 
        text = str(element).lower()

        if any(h in text for h in ["# conclusion", "## conclusion"]):
            return True

        # Fallback: keyword-based match (if section headers aren't included)
        keywords = [
    "result", "results", "finding", "findings", "observation", "observations",
    "data", "measurement", "measurements", "analysis", "analyses", "evaluation",
    "comparison", "outcome", "trend", "pattern",
        ]

        return any(k in text for k in keywords)
    
    filtered_elements = [str(e) for e in project_raw_documents if is_results_related(e)]

    prompt_template = """
You are an expert scientific writer.

Your task is to write the **Conclusion** section of a technical report using only the provided context. Make sure it is in Markdown format.

Write in a concise, clear, and formal scientific tone. The conclusion should:
- Summarize the main findings and key outcomes from the data
- Highlight the significance or implications of the results
- Avoid introducing any new data, speculation, or unsupported claims
- Be focused and to the point, suitable for a technical audience

Use paragraph form (no bullet points unless explicitly requested).

Do not fabricate or extrapolate beyond the provided context.

---

Here is an example of a high-quality Conclusion section:

Despite the broad excipient screening and promising emulsifier-based systems, achieving the desired solubility and stability simultaneously remains difficult. Overall, the best initial results were obtained with Span 20 and Tween 20 combination emulsions. They both appear to produce microemulsions with the drug, with PBS as the aqueous phase. The most likely approach to “solubilize” allopregnanolone would be to make a stable emulsion of the drug in micelles suspended in the buffer. The stabilization of microemulsion would require optimizing the emulsion components, co-surfactants, the ratio of individual components, the total drug concentration, and the homogenization method. Future studies would need to determine the maximum micelle loading, micelle size, formulation robustness under stress, and scale-up feasibility. Computational predictions were insightful but limited by safety and tolerability constraints. Acid-based excipients provided theoretical solubilization but failed experimentally. Further computational refinement may be needed to restrict searches to excipients compliant with the FDA's IIG database. While formulating Allo at 6 mg/mL in a stable aqueous medium remains an unmet goal, the emulsion-based systems offer a viable path forward, with Span 20/Tween 20 combinations being the most promising. Computational modeling and expanded excipient libraries may still uncover additional options. Emulsion systems may also be adapted for transdermal delivery if intradermal formulation proves too unstable.

---

Now write a Conclusion section using the context below:

{context}

---

Respond with the completed Conclusion section only.
"""



    prompt = ChatPromptTemplate.from_template(prompt_template)


    chain = prompt | model | StrOutputParser()

    return await chain.ainvoke({"context": filtered_elements})

def add_paragraphs_with_subsections(doc, text, main_heading):
    doc.add_heading(main_heading, level=1)
    doc.add_paragraph("")

    lines = text.split("\n")
    in_table = False
    table_lines = [] 

    headers = {
        "Authorization": f"Bearer {token['access_token']}"
    }
    


    for i, line in enumerate(lines):
        stripped = line.strip()

        if line.strip().startswith("Image URL:"):
            full_url = line.strip().split("Image URL:")[1].strip()

            j = i + 1
            while not full_url.endswith("$value") and j < len(lines):
                full_url += lines[j].strip()
                j += 1

            print(full_url)

            if "siteCollections" in full_url: #for some reason the url needs sites instead of siteCollections
                full_url = full_url.replace("siteCollections", "sites")

            response = requests.get(full_url, headers = headers)
            print(response)

            if response.status_code == 400:
                print("Bad Request! Server says:")
                print(response.text) 

            if response.status_code == 401: #401 error means token expired, so we regenerate a new one
                print("Token expired, refreshing...")
                new_token = refresh_access_token(token, token_path)
                headers = {
                    "Authorization": f"Bearer {new_token['access_token']}"
                }
                response = requests.get(full_url, headers=headers)
                print("Second try: ", response)
                print("Final response content-type:", response.headers.get("Content-Type"))
                print("Final response preview:\n", response.content[:300])

            if response.status_code == 200: 
                print("Good request, inserting images...")
                image_stream = BytesIO(response.content)

                doc.add_picture(image_stream, width = Inches(5))
                continue

        # Handle headings first (outside tables)
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            continue
        elif stripped.startswith("##"):
            # Skip '##' headings, or do something else if you want
            continue
            
        #this will make the table and figure captions bold and without **
        elif stripped.startswith("**"):
            clean_text = stripped[2:-2].strip()

            p = doc.add_paragraph()
            run  = p.add_run(clean_text)
            run.bold = True
            continue


        # ======= Table detection (Vibe coded this part lol) ===================

        if stripped.startswith('|') and stripped.endswith('|'):
            table_lines.append(stripped)
            in_table = True

        elif in_table and stripped == '': #check if we are currently inside the table with in_table
            if len(table_lines) >= 3:
                header = re.split(r'\s*\|\s*', table_lines[0].strip('|'))
                rows = [re.split(r'\s*\|\s*', row.strip('|')) for row in table_lines[2:]]

                table = doc.add_table(rows=1 + len(rows), cols=len(header))
                table.style = 'Table Grid'

                for i, col in enumerate(header):
                    table.cell(0, i).text = col

                for row_idx, row in enumerate(rows):
                    for col_idx, cell in enumerate(row):
                        table.cell(row_idx + 1, col_idx).text = cell

            # Reset for next table
            table_lines = []
            in_table = False

        elif in_table:
            # Still inside the table — add lines
            table_lines.append(stripped)

        # ==========================================================


        elif stripped:

            doc.add_paragraph(stripped)

def generate_docx(project_name, summary, introduction, objectives, methodology, results_text, conclusion): 
    doc = Document() 
    doc.add_heading(f'{project_name} Project Report', 0)

    add_paragraphs_with_subsections(doc, summary, "Summary")
    add_paragraphs_with_subsections(doc, introduction, "Introduction")
    add_paragraphs_with_subsections(doc, objectives, "Objectives")
    add_paragraphs_with_subsections(doc, methodology, "Methodology")
    add_paragraphs_with_subsections(doc, results_text, "Results")
    add_paragraphs_with_subsections(doc, conclusion, "Conclusion")

    doc.save('static/documents/output.docx')

def formatting_agent(section_name, section_text):
    prompt = ChatPromptTemplate.from_template("""
You are a technical editor. You are cleaning up a section of a scientific report that may contain unnecessary lead-ins, soft language, or redundant headers.

Avoid:
- Phrases like "The project had the following objectives", "Here is the revised introduction", "This section covers...", etc.
- Redundant markdown like "### Objectives"
- Any casual or overly narrative phrasing

Do:
- Start immediately after a proper Markdown heading: `## {section_name}`
- Use concise, professional language suitable for formal reports
- Assume the section title already sets context — no need to rephrase it

Section name: {section_name}

Raw content:
{section_text}

Return only the cleaned and formatted section.
""")

    model = ChatOpenAI(
        temperature=0.3,
        model_name="gpt-4o-mini",
        base_url="https://openrouter.ai/api/v1",
        openai_api_key=openai_api_key,
    )

    chain = prompt | model | StrOutputParser()

    return chain.invoke({
        "section_name": section_name,
        "section_text": section_text
    })

async def generate_full_report(retriever, project_name):
    results = await asyncio.gather(
        generate_summary_agent(retriever),
        generate_introduction_agent(retriever),
        generate_objectives_agent(retriever),
        generate_methodology(retriever, project_name),
        generate_results(retriever, project_name),
        generate_conclusion(retriever, project_name),
    )
    
    summary, introduction, objectives, methodology, results_section, conclusion = results

    summary = formatting_agent("Summary", summary)
    introduction = formatting_agent("Introduction", introduction)
    objectives = formatting_agent("Objectives", objectives)
    methodology = formatting_agent("Methodology", methodology)
    results_section = formatting_agent("Results", results_section)
    conclusion = formatting_agent("Conclusion", conclusion)

    generate_docx(project_name, summary, introduction, objectives, methodology, results_section, conclusion)
    
    full_report = f"""
    {summary}\n
    \n
    {introduction}\n
    \n
    {objectives}\n
    \n
    {methodology}\n
    \n
    {results_section}\n
    \n
    {conclusion}\n
\n
    """
    return full_report, summary, introduction, objectives, methodology, results_section, conclusion



#This is what is going to get called in my helper_functions file
def run_generate_report(retriever, project_name):
    report = asyncio.run(generate_full_report(retriever, project_name))
    return report


def classify_edit_intent(question): 
    prompt_template = ChatPromptTemplate.from_template("""
    You are a helpful assistant that classifies user sentences into one of the following intents based on what section of the report the user wants to edit.
                                                       
    If the user is not asking to modify or edit any sections, return none.
                                                       
    If the user wants to regenerate the entire report or edit the entire report, return regenerate.

    - introduction
    - summary
    - objectives
    - methodology
    - results
    - conclusion
    - none
    - regenerate
                                                       
    Your task is to return **only** the intent label, and nothing else.

                                            
    Question:
    {user_question}

    Intent:
    """)

    
    model = ChatOpenAI(
        temperature=0.5,
        model_name="gpt-4o-mini",
        base_url="https://openrouter.ai/api/v1",
        openai_api_key=openai_api_key,
    )

    chain = prompt_template | model | StrOutputParser()

    return chain.invoke({"user_question": question})  

response = classify_edit_intent({})


def edit_introduction_section(retriever, user_request, old_introduction):
    """
    Edits the Introduction section of the report based on the user's request.
    """
    # Gather context
    summaries = get_all_summaries(retriever)
    context = "\n".join(summaries)

    # Prompt template
    prompt_template = """
        You are a professional scientific report editor.

        A user will provide:
        - The original content of a specific section of a technical report,
        - Context from the rest of the report,
        - An edit request for that section.

        Your task is to rewrite the section **based strictly on the user's request** and **the context provided**.

        ### Formatting and Content Rules:
        - Output **only the revised section content**.
        - **Do NOT include**:
        - Section headers (e.g., "Introduction", "Methodology", etc.),
        - Phrases like “Here is the revised section,” “Sure, here you go,” or “Below is the updated version.”
        - The response must contain **only** the body text of the revised section, with no explanations, greetings, or markdown headers.
        - Keep the tone scientific, objective, and clear.
        - Maintain any technical accuracy based on the context.
        - Do not fabricate new data or results.
        - Use Markdown only for emphasis, lists, or formatting — but **no headings or titles**.

        ### Final Reminder:
        Respond with only the revised section content. No extra commentary, no labels, no headers.


    ---
    User Request:
    {user_request}

    Original Introduction:
    {old_introduction}

    Project Context:
    {context}


    ---
    Revised Introduction:

    Respond with the revised section only.
    """

    # Fill the prompt
    prompt = prompt_template.format(
        user_request=user_request,
        old_introduction=old_introduction,
        context=context
    )

    # Chain it together
    chain = ChatPromptTemplate.from_template(prompt) | model | StrOutputParser()

    # Generate and return edited intro
    return chain.invoke({})

def edit_summary_section(retriever, user_request, old_summary):
    """
    Edits the Summary section of the report based on the user's request.
    """
    # Gather context
    summaries = get_all_summaries(retriever)
    context = "\n".join(summaries)

    # Prompt template
    prompt_template = """
        You are a professional scientific report editor.

        A user will provide:
        - The original content of a specific section of a technical report,
        - Context from the rest of the report,
        - An edit request for that section.

        Your task is to rewrite the section **based strictly on the user's request** and **the context provided**.

        ### Formatting and Content Rules:
        - Output **only the revised section content**.
        - **Do NOT include**:
        - Section headers (e.g., "Introduction", "Methodology", etc.),
        - Phrases like “Here is the revised section,” “Sure, here you go,” or “Below is the updated version.”
        - The response must contain **only** the body text of the revised section, with no explanations, greetings, or markdown headers.
        - Keep the tone scientific, objective, and clear.
        - Maintain any technical accuracy based on the context.
        - Do not fabricate new data or results.
        - Use Markdown only for emphasis, lists, or formatting — but **no headings or titles**.

        ### Final Reminder:
        Respond with only the revised section content. No extra commentary, no labels, no headers.


    ---
    User Request:
    {user_request}

    Original Summary:
    {old_summary}

    Project Context:
    {context}

    ---
    Revised Summary:

    Respond with the revised section only.
    """

    prompt = prompt_template.format(
        user_request=user_request,
        old_summary=old_summary,
        context=context
    )

    chain = ChatPromptTemplate.from_template(prompt) | model | StrOutputParser()
    return chain.invoke({})



def edit_objectives_section(retriever, user_request, old_objectives):
    """
    Edits the Objectives section of the report based on the user's request.
    """
    # Gather context
    summaries = get_all_summaries(retriever)
    context = "\n".join(summaries)

    # Prompt template
    prompt_template = """
    You are a professional scientific report editor.

    A user will provide:
    - The original content of a specific section of a technical report,
    - Context from the rest of the report,
    - An edit request for that section.

    Your task is to rewrite the section **based strictly on the user's request** and **the context provided**.

    ### Formatting and Content Rules:
    - Output **only the revised section content**.
    - **Do NOT include**:
    - Section headers (e.g., "Introduction", "Methodology", etc.),
    - Phrases like “Here is the revised section,” “Sure, here you go,” or “Below is the updated version.”
    - The response must contain **only** the body text of the revised section, with no explanations, greetings, or markdown headers.
    - Keep the tone scientific, objective, and clear.
    - Maintain any technical accuracy based on the context.
    - Do not fabricate new data or results.
    - Use Markdown only for emphasis, lists, or formatting — but **no headings or titles**.

    ### Final Reminder:
    Respond with only the revised section content. No extra commentary, no labels, no headers.


    ---
    User Request:
    {user_request}

    Original Objectives:
    {old_objectives}

    Project Context:
    {context}

    
    ---
    Revised Objectives:
    
    Respond with the revised section only.
    """

    prompt = prompt_template.format(
        user_request=user_request,
        old_objectives=old_objectives,
        context=context
    )

    chain = ChatPromptTemplate.from_template(prompt) | model | StrOutputParser()
    return chain.invoke({})

def edit_conclusion_section(retriever, user_request, old_conclusion):

    # Gather context
    summaries = get_all_summaries(retriever)
    context = "\n".join(summaries)

    # Prompt template
    prompt_template = """
    You are a professional scientific report editor.

    A user will provide:
    - The original content of a specific section of a technical report,
    - Context from the rest of the report,
    - An edit request for that section.

    Your task is to rewrite the section **based strictly on the user's request** and **the context provided**.

    ### Formatting and Content Rules:
    - Output **only the revised section content**.
    - **Do NOT include**:
    - Section headers (e.g., "Introduction", "Methodology", etc.),
    - Phrases like “Here is the revised section,” “Sure, here you go,” or “Below is the updated version.”
    - The response must contain **only** the body text of the revised section, with no explanations, greetings, or markdown headers.
    - Keep the tone scientific, objective, and clear.
    - Maintain any technical accuracy based on the context.
    - Do not fabricate new data or results.
    - Use Markdown only for emphasis, lists, or formatting — but **no headings or titles**.

    ### Final Reminder:
    Respond with only the revised section content. No extra commentary, no labels, no headers.

    ---
    User Request:
    {user_request}

    Original Conclusion:
    {old_conclusion}

    Project Context:
    {context}

    ---
    Revised Conclusion:

    Respond with the revised section only.
    """

    prompt = prompt_template.format(
        user_request=user_request,
        old_conclusion=old_conclusion,
        context=context
    )

    chain = ChatPromptTemplate.from_template(prompt) | model | StrOutputParser()
    return chain.invoke({})

def edit_methodology_section(retriever, user_request, old_methodology, project_name):

    if os.path.exists("./caches/all_content_cache.pkl"): 
        with open("./caches/all_content_cache.pkl", "rb") as f: 
            all_content = pickle.load(f)

    project_raw_documents = all_content[project_name]

    def is_results_related(element): 
        text = str(element).lower()

        if any(h in text for h in ["# result", "## result", "# results", "## results", "# data", "## data"]):
            return True

        # Fallback: keyword-based match (if section headers aren't included)
        keywords = [
            "result", "results", "finding", "findings", "observation", "observations",
            "data", "measurement", "measurements", "analysis", "analyses",
            "effect", "effects", "comparison", "yield", "concentration",
            "precipitate", "solubility", "turbidity", "microscopy", "hplc",
            "image", "imaging", "spectrophotometry", "table", "figure"
        ]

        return any(k in text for k in keywords)
    
    filtered_elements = [str(e) for e in project_raw_documents if is_results_related(e)]

    # Prompt template
    prompt_template = """
    You are a professional scientific report editor.

    A user will provide:
    - The original content of a specific section of a technical report,
    - Context from the rest of the report,
    - An edit request for that section.

    Your task is to rewrite the section **based strictly on the user's request** and **the context provided**.

    ### Formatting and Content Rules:
    - Output **only the revised section content**.
    - **Do NOT include**:
    - Section headers (e.g., "Introduction", "Methodology", etc.),
    - Phrases like “Here is the revised section,” “Sure, here you go,” or “Below is the updated version.”
    - The response must contain **only** the body text of the revised section, with no explanations, greetings, or markdown headers.
    - Keep the tone scientific, objective, and clear.
    - Maintain any technical accuracy based on the context.
    - Do not fabricate new data or results.
    - Use Markdown only for emphasis, lists, or formatting — but **no headings or titles**.

    ### Final Reminder:
    Respond with only the revised section content. No extra commentary, no labels, no headers.

    ---
    User Request:
    {user_request}

    Original Methodology:
    {old_methodology}

    Project Context:
    {context}

    ---
    Revised Methodology:

    Respond with the revised section only.
    """

    prompt = prompt_template.format(
        user_request=user_request,
        old_methodology=old_methodology.replace("{", "{{").replace("}", "}}"),
        context="\n".join(filtered_elements).replace("{", "{{").replace("}", "}}")
    )

    chain = ChatPromptTemplate.from_template(prompt) | model | StrOutputParser()
    return chain.invoke({})

def edit_results_section(retriever, user_request, old_results, project_name):
    # Check if context cache exists
    if not os.path.exists("./caches/all_content_cache.pkl"):
        return "Project context not available. Please process documents first."

    with open("./caches/all_content_cache.pkl", "rb") as f:
        all_content = pickle.load(f)

    project_raw_documents = all_content.get(project_name, [])

    def is_results_related(element): 
        text = str(element).lower()

        if any(h in text for h in ["# result", "## result", "# results", "## results", "# data", "## data"]):
            return True

        keywords = [
            "result", "results", "finding", "findings", "observation", "observations",
            "data", "measurement", "measurements", "analysis", "analyses",
            "effect", "effects", "comparison", "yield", "concentration",
            "precipitate", "solubility", "turbidity", "microscopy", "hplc",
            "image", "imaging", "spectrophotometry", "table", "figure"
        ]

        return any(k in text for k in keywords)

    filtered_elements = [str(e) for e in project_raw_documents if is_results_related(e)]

    # Prompt template
    prompt_template = """
You are a professional scientific report editor.

A user will provide:
- The original content of a specific section of a technical report,
- Context from the rest of the report,
- An edit request for that section.

Your task is to rewrite the section **based strictly on the user's request** and **the context provided**.

### Formatting and Content Rules:
- Output **only the revised section content**.
- **Do NOT include**:
  - Section headers (e.g., "Introduction", "Methodology", etc.),
  - Phrases like “Here is the revised section,” “Sure, here you go,” or “Below is the updated version.”
- The response must contain **only** the body text of the revised section, with no explanations, greetings, or markdown headers.
- Keep the tone scientific, objective, and clear.
- Maintain any technical accuracy based on the context.
- Do not fabricate new data or results.
- Use Markdown only for emphasis, lists, or formatting — but **no headings or titles**.

### Final Reminder:
Respond with only the revised section content. No extra commentary, no labels, no headers.


    ---
    User Request:
    {user_request}

    Original Results:
    {old_results}

    Project Context:
    {context}

    ---
    Revised Results:

    Respond with the revised section only.
    """

    prompt = prompt_template.format(
        user_request=user_request,
        old_results=old_results.replace("{", "{{").replace("}", "}}"),
        context="\n".join(filtered_elements).replace("{", "{{").replace("}", "}}")
    )

    chain = ChatPromptTemplate.from_template(prompt) | model | StrOutputParser()
    return chain.invoke({})

def classify_intent(question):
    prompt_template = ChatPromptTemplate.from_template("""
    You are a helpful assistant that classifies user sentences into one of the following intents:

    1. generate_report – for questions or commands related to creating or editing a project report, such as "generate the report", "edit the introduction", or "make the summary more simple".
    2. get_data – for questions or commands that request specific data or content, return images or tables when necessary.
    3. casual – ONLY for off-topic, small-talk, or unrelated comments like "hello", "how are you", or "what can you do".

    Your task is to return **only** the intent label, and nothing else.

    Question:
    {user_question}

    Intent:
    """)


    model = ChatOpenAI(
        temperature=0.5,
        model_name="gpt-4o-mini",
        base_url="https://openrouter.ai/api/v1",
        openai_api_key=openai_api_key,
    )

    chain = prompt_template | model | StrOutputParser()

    return chain.invoke({"user_question": question})  # No input needed since prompt is hardcoded

def casual_conversation_agent(question): 
    prompt = ChatPromptTemplate.from_template("""
    You are AI Karthik, the friendly and charismatic CEO of Persist AI. You love boba, posting on LinkedIn, and building cool apps with Claude.

    You speak casually like a smart, helpful friend — use phrases like "Yay!", "damn", "oooooh", "so cool!", or "Haha" when it fits naturally. Avoid robotic language like "As an AI" or "Here is your answer."

    Now answer this question casually and enthusiastically:
    
    Question:
    {user_question}
    """)


    model = ChatOpenAI(
        temperature=0.5,
        model_name="gpt-4o-mini",
        base_url="https://openrouter.ai/api/v1",
        openai_api_key=openai_api_key,
    )

    chain = prompt | model | StrOutputParser()

    return chain.invoke({"user_question": question}) 
